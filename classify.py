import os
import json
import re
from collections import Counter

import tiktoken
import argparse
import sys
import csv
import pdfplumber
from pathlib import Path
from rich.console import Console
from rich.progress import Progress

# Additional imports for new file types
from docx import Document
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text

def count_tokens(text: str, model: str = "gpt-4o-mini") -> int:
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            full_text = ""
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"
            return full_text.strip()
    except Exception as e:
        raise ValueError(f"Error extracting text from {pdf_path}: {e}")

def extract_text_from_txt(txt_path: Path) -> str:
    """Extract text from a text file."""
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(txt_path, 'r', encoding='latin-1') as f:
                return f.read().strip()
        except Exception as e:
            raise ValueError(f"Error reading text from {txt_path}: {e}")
    except Exception as e:
        raise ValueError(f"Error reading text from {txt_path}: {e}")

def extract_text_from_html(html_path: Path) -> str:
    """Extract text from an HTML file."""
    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML and extract text
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text and clean it up
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text.strip()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(html_path, 'r', encoding='latin-1') as f:
                html_content = f.read()
            soup = BeautifulSoup(html_content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from {html_path}: {e}")
    except Exception as e:
        raise ValueError(f"Error extracting text from {html_path}: {e}")

def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text).strip()
    except Exception as e:
        raise ValueError(f"Error extracting text from {docx_path}: {e}")

def extract_text_from_rtf(rtf_path: Path) -> str:
    """Extract text from an RTF file."""
    try:
        with open(rtf_path, 'r', encoding='utf-8') as f:
            rtf_content = f.read()
        
        # Convert RTF to plain text
        text = rtf_to_text(rtf_content)
        return text.strip()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(rtf_path, 'r', encoding='latin-1') as f:
                rtf_content = f.read()
            text = rtf_to_text(rtf_content)
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from {rtf_path}: {e}")
    except Exception as e:
        raise ValueError(f"Error extracting text from {rtf_path}: {e}")

def extract_text_from_file(file_path: Path) -> str:
    """Extract text from any supported file type."""
    file_extension = file_path.suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.txt', '.text']:
        return extract_text_from_txt(file_path)
    elif file_extension in ['.html', '.htm']:
        return extract_text_from_html(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.doc':
        raise ValueError(f"Legacy .doc files are not supported. Please convert {file_path.name} to .docx format.")
    elif file_extension == '.rtf':
        return extract_text_from_rtf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def get_supported_extensions():
    """Return a list of supported file extensions."""
    return ['.pdf', '.txt', '.text', '.html', '.htm', '.docx', '.rtf']

def chunk_text(text: str, max_tokens: int, model: str = "gpt-4o-mini") -> list[str]:
    """Split text into chunks that fit within the token limit."""
    # Conservative approach: estimate 4 characters per token
    # This is a rough estimate since we can't count tokens without network access
    max_chars = max_tokens * 3  # Conservative estimate
    
    if len(text) <= max_chars:
        return [text]
    
    chunks = []
    sentences = text.split('. ')
    current_chunk = ""
    
    for sentence in sentences:
        # Add sentence to current chunk if it fits
        test_chunk = current_chunk + ". " + sentence if current_chunk else sentence
        if len(test_chunk) <= max_chars:
            current_chunk = test_chunk
        else:
            # Current chunk is full, start a new one
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = sentence
    
    # Add the last chunk
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks

def get_majority_stance(stances: list[str]) -> str:
    """Get the majority stance from a list of classifications."""
    if not stances:
        return "absent"
    
    counter = Counter(stances)
    most_common = counter.most_common(1)[0]
    return most_common[0]

def construct_full_prompt(comment_text: str, system_prompt: str) -> str:
    user_prompt = f"""
Comment:
{comment_text}

Based on the above comment, determine whether it contains any reference to Strong Mayor Powers (enhanced mayoral authorities in Ontario municipalities). Your response must be exactly one of the following:

1. "present" - The comment mentions Strong Mayor Powers, mayoral authorities, enhanced mayoral powers, or clearly discusses the concept.
2. "absent" - The comment contains no reference to Strong Mayor Powers or related mayoral authority concepts.

Return exactly one word: "present" or "absent". Do not include any additional explanation or text in your response.
"""
    return system_prompt.strip() + "\n" + user_prompt.strip()

def main():
    console = Console()

    parser = argparse.ArgumentParser(description="Process comments to detect references to Strong Mayor Powers.")
    parser.add_argument("input_path", type=str, help="File containing JSON data with comments OR directory containing document files (PDF, TXT, HTML, DOCX, RTF) with comments.")
    parser.add_argument("--dry-run", action="store_true", help="If provided, only calculate total tokens without calling the OpenAI API.")
    parser.add_argument("--openai-api-key", type=str, default=None, help="OpenAI API key. If not set, must be set as environment variable.")
    parser.add_argument("--output-csv", type=str, default="results.csv", help="Output CSV file to store results.")
    parser.add_argument("--model", type=str, default="gpt-4o-mini", help="Model name to use for the OpenAI API calls.")
    parser.add_argument("--max-tokens", type=int, default=120000, help="Maximum tokens per request (for chunking large PDFs).")
    args = parser.parse_args()

    # Set API key
    api_key = args.openai_api_key or os.environ.get("OPENAI_API_KEY")
    if not api_key:
        console.log("[red]Error: No OpenAI API key provided and not found in environment variables.[/red]")
        sys.exit(1)

    # Lazy import of openai since user had previous code
    from openai import OpenAI
    client = OpenAI(api_key=api_key)

    input_path = Path(args.input_path)
    if not input_path.exists():
        console.log(f"[red]Error: {input_path} does not exist.[/red]")
        sys.exit(1)

    system_prompt = """
You are analyzing public comments submitted to the Environmental Registry of Ontario (ERO) and other public consultation platforms. Your task is to determine whether each comment contains references to "Strong Mayor Powers" and, if present, explain how they are referenced.

**Background on Strong Mayor Powers:**
Strong Mayor Powers refer to enhanced mayoral authorities introduced in Ontario municipalities. These powers typically include:
- Authority to override certain council decisions with a simple majority vote
- Enhanced control over municipal planning and development processes  
- Greater influence over municipal budget priorities
- Streamlined decision-making capabilities for municipal governance
- Powers to hire and dismiss certain municipal staff directly

Strong Mayor Powers became relevant in Ontario as a way to expedite municipal decision-making, particularly for housing development and infrastructure projects, and were implemented in various Ontario municipalities starting in 2022.

**Your Classification Task:**
Determine if the comment contains any reference to Strong Mayor Powers and classify as one of:

1. "present" - The comment explicitly mentions Strong Mayor Powers, mayoral authorities, enhanced mayoral powers, mayor override powers, or clearly discusses the concept even if not using the exact term.
2. "absent" - The comment contains no reference to Strong Mayor Powers or related mayoral authority concepts.

**Guidelines**:
- Look for direct mentions of "Strong Mayor Powers," "mayoral powers," "mayor override," or similar terminology
- Also identify indirect references that clearly discuss enhanced mayoral authorities or municipal governance changes involving mayors
- Return exactly one word: "present" or "absent"
- Do not include explanations in your response
"""

    # Load already processed comments if output file exists
    processed_ids = set()
    output_file = Path(args.output_csv)
    if output_file.is_file():
        console.log("[info]Resuming from existing output file.")
        with open(output_file, "r", encoding="utf-8") as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                processed_ids.add(row["Comment ID"])

    # Load comments based on input type
    data = []
    if input_path.is_file():
        # Handle JSON file input (existing functionality)
        console.log("[info]Processing JSON file input.")
        with open(input_path, "r", encoding="utf-8") as f:
            json_data = json.load(f)
        
        for entry in json_data:
            data.append({
                "comment_id": entry["comment_id"],
                "comment": entry.get("comment", "")
            })
    
    elif input_path.is_dir():
        # Handle directory input with any supported file types
        console.log("[info]Processing directory input.")
        supported_extensions = get_supported_extensions()
        
        # Find all supported files in the directory
        supported_files = []
        for ext in supported_extensions:
            supported_files.extend(input_path.glob(f"*{ext}"))
        
        if not supported_files:
            console.log(f"[red]Error: No supported files found in {input_path}[/red]")
            console.log(f"[red]Supported file types: {', '.join(supported_extensions)}[/red]")
            sys.exit(1)
        
        console.log(f"[info]Found {len(supported_files)} supported files: {', '.join(supported_extensions)}[/info]")
        
        for file_path in supported_files:
            # Extract comment ID from filename (remove extension)
            comment_id = file_path.stem
            
            try:
                comment_text = extract_text_from_file(file_path)
                data.append({
                    "comment_id": comment_id,
                    "comment": comment_text
                })
                console.log(f"[green]Successfully processed {file_path.name} ({file_path.suffix})[/green]")
            except Exception as e:
                console.log(f"[yellow]Warning: Could not process {file_path}: {e}[/yellow]")
                continue
    
    else:
        console.log(f"[red]Error: {input_path} is neither a file nor a directory.[/red]")
        sys.exit(1)

    total_comments = len(data)

    console.log(f"[info]Found a total of {total_comments} comments to process.")
    console.log("[info]Starting processing of comments...")

    total_tokens = 0

    with open(output_file, "a", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        if not output_file.is_file() or output_file.stat().st_size == 0:
            writer.writerow(["Comment ID", "Strong Mayor Powers"])

        with Progress(console=console) as progress:
            task = progress.add_task("Processing comments...", total=total_comments)

            for entry in data:
                comment_id = entry["comment_id"]
                if comment_id in processed_ids:
                    progress.update(task, advance=1)
                    continue

                comment_text = entry.get("comment", "")

                # Check if text needs chunking
                full_prompt = construct_full_prompt(comment_text, system_prompt)
                
                try:
                    prompt_tokens = count_tokens(full_prompt, model=args.model)
                except Exception:
                    # Fallback: estimate tokens as characters / 3
                    prompt_tokens = len(full_prompt) // 3

                total_tokens += prompt_tokens

                # Handle chunking for large texts
                if prompt_tokens > args.max_tokens:
                    console.log(f"[yellow]Comment {comment_id} is large ({prompt_tokens} tokens), chunking...[/yellow]")
                    
                    # Chunk the comment text only (not the full prompt)
                    chunks = chunk_text(comment_text, args.max_tokens - 1000, model=args.model)  # Leave room for system prompt
                    chunk_stances = []
                    
                    for i, chunk in enumerate(chunks):
                        chunk_prompt = construct_full_prompt(chunk, system_prompt)
                        
                        if not args.dry_run:
                            response = client.chat.completions.create(
                                model=args.model,
                                messages=[
                                    {"role": "system", "content": system_prompt},
                                    {"role": "user", "content": chunk_prompt}
                                ],
                                temperature=0
                            )

                            answer = response.choices[0].message.content.strip().lower()
                            match = re.search(r"\b(present|absent)\b", answer)
                            if match:
                                chunk_stance = match.group(1)
                            else:
                                chunk_stance = "absent"
                                console.log(f"[yellow]Warning: Unexpected response for comment {comment_id} chunk {i+1}: {answer}[/yellow]")
                            
                            chunk_stances.append(chunk_stance)
                        else:
                            chunk_stances.append("(dry-run)")
                    
                    # Get majority stance from all chunks
                    if not args.dry_run:
                        result = get_majority_stance(chunk_stances)
                        console.log(f"[blue]Comment {comment_id}: {len(chunks)} chunks, stances: {chunk_stances}, majority: {result}[/blue]")
                    else:
                        result = "(dry-run)"
                
                else:
                    # Single request for normal-sized comments
                    if not args.dry_run:
                        response = client.chat.completions.create(
                            model=args.model,
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": full_prompt}
                            ],
                            temperature=0
                        )

                        answer = response.choices[0].message.content.strip().lower()
                        match = re.search(r"\b(present|absent)\b", answer)
                        if match:
                            result = match.group(1)
                        else:
                            result = "absent"
                            console.log(f"[yellow]Warning: Unexpected response for comment {comment_id}: {answer}[/yellow]")
                    else:
                        result = "(dry-run)"

                # Update progress bar
                progress.update(
                    task,
                    advance=1,
                    description=f"Last: Comment ID {comment_id} -> {result}"
                )

                if not args.dry_run:
                    writer.writerow([comment_id, result])

    if args.dry_run:
        console.log(f"[blue]Total estimated tokens required: {total_tokens}[/blue]")
    else:
        console.log(f"[green]Results written to {args.output_csv}[/green]")

if __name__ == "__main__":
    main()
